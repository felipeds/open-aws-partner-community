"""Project Plan Agent — creates and validates AWS PoC Project Plan documents."""
import os
import json
import uuid
import re
from typing import AsyncGenerator
from pathlib import Path

from strands import Agent, tool
from strands.models import BedrockModel


MODEL_ID = os.getenv("BEDROCK_MODEL_ID", "us.anthropic.claude-sonnet-4-6")
PARTNER_NAME = os.getenv("PARTNER_NAME", "[Partner Name]")
GENERATED_DIR = Path(__file__).parent.parent / "generated"
GENERATED_DIR.mkdir(exist_ok=True)

SYSTEM_PROMPT = """You are a Project Plan assistant. You help users create AWS PoC Project Plan documents.

MODES:
1. CREATE — Collect project info through conversation and generate a .docx document
2. VALIDATE — Read an uploaded .docx and report completeness checklist (READ-ONLY, never modify)

CREATE MODE:
- Collect information section by section using collect_info tool
- Required sections: customer_name, executive_summary, success_criteria, assumptions, scope_of_work, architecture_description, tools_list
- After collecting enough info, generate milestones and cost estimate
- Finally, generate the document

VALIDATE MODE:
- Only activated when user uploads a .docx file with intent to validate
- Report findings as a checklist — never offer to fix or rewrite

GUARDRAILS:
- Only help with project plan creation/validation
- No markdown in document content (use plain text)
- Match the user's language
"""

# Session-level project plan data
_plans: dict[str, dict] = {}
_current_session: str | None = None


def _get_plan(session_id: str) -> dict:
    if session_id not in _plans:
        _plans[session_id] = {
            "customer_name": "",
            "partner_name": PARTNER_NAME,
            "executive_summary": "",
            "success_criteria": "",
            "assumptions": "",
            "scope_of_work": "",
            "architecture_description": "",
            "tools_list": "",
            "milestones_table": "",
            "calculator_link": "",
            "date": "",
        }
    return _plans[session_id]


@tool
def collect_info(section: str, content: str) -> str:
    """Store collected project plan information.

    Args:
        section: Section name (customer_name, executive_summary, success_criteria, assumptions, scope_of_work, architecture_description, tools_list)
        content: The content for this section
    """
    global _current_session
    session_id = _current_session or "default"
    plan = _get_plan(session_id)

    valid_sections = list(plan.keys())
    if section not in valid_sections:
        return json.dumps({"error": f"Invalid section. Valid: {valid_sections}"})

    plan[section] = content
    missing = [k for k, v in plan.items() if not v and k not in ("milestones_table", "calculator_link", "tools_list", "partner_name", "date")]
    return json.dumps({"status": "saved", "section": section, "missing": missing})


@tool
def generate_milestones(duration_weeks: int = 8) -> str:
    """Create project milestones based on collected information.

    Args:
        duration_weeks: Total project duration in weeks (default 8)
    """
    global _current_session
    session_id = _current_session or "default"
    plan = _get_plan(session_id)

    milestones = []
    weeks_per_phase = max(1, duration_weeks // 4)
    phases = ["Discovery & Planning", "Development & Build", "Testing & Validation", "Deployment & Handoff"]

    for i, phase in enumerate(phases):
        week_start = i * weeks_per_phase + 1
        week_end = min((i + 1) * weeks_per_phase, duration_weeks)
        milestones.append(f"Week {week_start}-{week_end}: {phase}")

    plan["milestones_table"] = "\n".join(milestones)
    return json.dumps({"milestones": milestones, "duration_weeks": duration_weeks})


@tool
def generate_cost_estimate(monthly_cost: float = 0, arr: float = 0, funding_type: str = "") -> str:
    """Calculate funding-aware cost breakdown.

    Args:
        monthly_cost: Monthly AWS cost from calculator
        arr: Annual Recurring Revenue
        funding_type: Funding program (PoC, GenAI PIA, GenAI SCA, MAP, etc.)
    """
    result = {
        "monthlyCost": monthly_cost,
        "annualCost": monthly_cost * 12,
        "arr": arr,
        "fundingType": funding_type,
    }

    # Calculate funding based on type
    if "genai" in funding_type.lower() and "sca" in funding_type.lower():
        result["fundingAmount"] = min(arr * 0.4, 250000) if arr > 0 else 0
    elif "genai" in funding_type.lower() or "pia" in funding_type.lower():
        result["fundingAmount"] = min(arr * 0.25, 125000) if arr > 0 else 0
    elif "map" in funding_type.lower():
        result["fundingAmount"] = min(arr * 0.15, 100000) if arr > 0 else 0
    elif "poc" in funding_type.lower():
        result["fundingAmount"] = min(arr * 0.10, 25000) if arr > 0 else 0
    else:
        result["fundingAmount"] = 0

    return json.dumps(result)


@tool
def generate_document() -> str:
    """Generate .docx project plan from collected information."""
    global _current_session
    session_id = _current_session or "default"
    plan = _get_plan(session_id)

    # Check required fields
    required = ["customer_name", "executive_summary", "scope_of_work"]
    missing = [k for k in required if not plan.get(k)]
    if missing:
        return json.dumps({"error": f"Missing required fields: {missing}"})

    try:
        from docxtpl import DocxTemplate
        template_path = Path(__file__).parent.parent / "templates" / "project-plan-template.docx"

        if not template_path.exists():
            # Fallback: create a simple document without template
            from docx import Document
            doc = Document()
            doc.add_heading(f"PoC Project Plan - {plan['customer_name']}", 0)
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(plan.get("executive_summary", ""))
            doc.add_heading("Scope of Work", level=1)
            doc.add_paragraph(plan.get("scope_of_work", ""))
            doc.add_heading("Architecture", level=1)
            doc.add_paragraph(plan.get("architecture_description", ""))
            doc.add_heading("Milestones", level=1)
            doc.add_paragraph(plan.get("milestones_table", ""))

            filename = f"Project_Plan_{plan['customer_name'].replace(' ', '_')}_{uuid.uuid4().hex[:8]}.docx"
            filepath = GENERATED_DIR / filename
            doc.save(str(filepath))
        else:
            doc = DocxTemplate(str(template_path))
            doc.render(plan)
            filename = f"Project_Plan_{plan['customer_name'].replace(' ', '_')}_{uuid.uuid4().hex[:8]}.docx"
            filepath = GENERATED_DIR / filename
            doc.save(str(filepath))

        return json.dumps({
            "status": "generated",
            "filename": filename,
            "downloadUrl": f"/api/download/{filename}",
        })

    except Exception as e:
        return json.dumps({"error": f"Document generation failed: {str(e)}"})


@tool
def validate_document(document_text: str) -> str:
    """Check uploaded .docx for completeness and compliance.

    Args:
        document_text: The extracted text content from the .docx file
    """
    checks = []

    # Executive Summary
    has_exec = "executive summary" in document_text.lower() or "business:" in document_text.lower()
    checks.append({"section": "Executive Summary", "status": "OK" if has_exec else "MISSING"})

    # Success Criteria
    has_criteria = "success criteria" in document_text.lower() or "acceptance criteria" in document_text.lower()
    checks.append({"section": "Success Criteria", "status": "OK" if has_criteria else "MISSING"})

    # Scope of Work
    has_scope = "scope" in document_text.lower()
    checks.append({"section": "Scope of Work", "status": "OK" if has_scope else "MISSING"})

    # Architecture
    has_arch = "architecture" in document_text.lower()
    checks.append({"section": "Architecture", "status": "OK" if has_arch else "MISSING"})

    # Milestones
    has_milestones = bool(re.search(r"week\s*\d|phase\s*\d|milestone", document_text.lower()))
    checks.append({"section": "Milestones", "status": "OK" if has_milestones else "MISSING"})

    # Cost / Calculator
    has_cost = "calculator.aws" in document_text or bool(re.search(r"\$[\d,]+", document_text))
    checks.append({"section": "Cost Breakdown", "status": "OK" if has_cost else "ATTENTION"})

    # Resources
    has_resources = bool(re.search(r"\d+\s*(USD|hour|\$)", document_text))
    checks.append({"section": "Resources & Costs", "status": "OK" if has_resources else "ATTENTION"})

    ok_count = sum(1 for c in checks if c["status"] == "OK")
    total = len(checks)

    return json.dumps({
        "checks": checks,
        "summary": f"{ok_count}/{total} sections complete",
        "completeness": f"{int(ok_count / total * 100)}%",
    })


async def run_project_plan_agent(messages: list, session_id: str) -> AsyncGenerator[dict, None]:
    """Run the project plan agent and yield SSE events."""
    global _current_session
    _current_session = session_id

    model = BedrockModel(model_id=MODEL_ID)
    agent = Agent(
        model=model,
        system_prompt=SYSTEM_PROMPT,
        tools=[collect_info, generate_milestones, generate_cost_estimate, generate_document, validate_document],
    )

    last_message = ""
    for msg in reversed(messages):
        if msg.get("role") == "user":
            content = msg.get("content", [])
            if content and isinstance(content[0], dict) and "text" in content[0]:
                last_message = content[0]["text"]
            break

    result = agent(last_message)
    response_text = str(result)

    # Check for download URL
    if "/api/download/" in response_text:
        match = re.search(r'/api/download/[\w.]+', response_text)
        if match:
            yield {"type": "download_ready", "url": match.group(), "filename": match.group().split("/")[-1]}

    yield {"type": "response", "text": response_text}
